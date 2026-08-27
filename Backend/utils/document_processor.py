import fitz  # PyMuPDF
import sys
import os
import io
import re
import csv
import zipfile
import xml.etree.ElementTree as ET
import asyncio
import logging
from typing import Tuple, Optional
from concurrent.futures import ThreadPoolExecutor

try:
    import pytesseract
    if sys.platform.startswith('win'):
        default_tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        if os.path.exists(default_tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = default_tesseract_path
except ImportError:
    pytesseract = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import docx
except ImportError:
    docx = None

logger = logging.getLogger(__name__)


class DocumentProcessor:

    # -------------------------------------------------------------------------
    # 1. PDF EXTRACTION
    # -------------------------------------------------------------------------
    @staticmethod
    async def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF using PyMuPDF with OCR fallback for scanned pages"""
        try:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, DocumentProcessor._process_pdf_sync, file_content)
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed, trying fallback: {e}")
            return DocumentProcessor._extract_printable_text(file_content)

    @staticmethod
    def _process_pdf_sync(file_content: bytes) -> str:
        doc = fitz.Document(stream=file_content, filetype="pdf")
        text_parts = []
        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            # If page has very little text and pytesseract/PIL are available, run OCR
            if len(page_text.strip()) < 50 and pytesseract and Image:
                try:
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    image = Image.open(io.BytesIO(img_data))
                    ocr_text = pytesseract.image_to_string(image)
                    if ocr_text.strip():
                        page_text = f"{page_text}\n{ocr_text}"
                except Exception as ocr_err:
                    logger.debug(f"Page {page_num} OCR fallback failed: {ocr_err}")
            
            if page_text.strip():
                text_parts.append(f"--- Page {page_num} ---\n{page_text.strip()}")
        
        doc.close()
        return "\n\n".join(text_parts).strip()

    # -------------------------------------------------------------------------
    # 2. WORD DOCUMENTS (.docx, .doc)
    # -------------------------------------------------------------------------
    @staticmethod
    async def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX/DOC file"""
        try:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, DocumentProcessor._process_docx_sync, file_content)
        except Exception as e:
            logger.warning(f"DOCX extraction fallback: {e}")
            return DocumentProcessor._extract_printable_text(file_content)

    @staticmethod
    def _process_docx_sync(file_content: bytes) -> str:
        # Try python-docx first
        if docx:
            try:
                bio = io.BytesIO(file_content)
                doc = docx.Document(bio)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # Also extract table text
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            tables_text.append(" | ".join(row_text))
                combined = paragraphs + tables_text
                if combined:
                    return "\n\n".join(combined).strip()
            except Exception as docx_err:
                logger.debug(f"python-docx parsing failed, trying raw XML: {docx_err}")

        # Raw OpenXML Zip fallback for .docx
        try:
            with zipfile.ZipFile(io.BytesIO(file_content)) as z:
                if 'word/document.xml' in z.namelist():
                    xml_content = z.read('word/document.xml')
                    tree = ET.fromstring(xml_content)
                    # Extract all text nodes in document
                    texts = [elem.text for elem in tree.iter() if elem.text and elem.text.strip()]
                    if texts:
                        return "\n".join(texts).strip()
        except Exception:
            pass

        # Fallback for binary .doc or damaged files
        return DocumentProcessor._extract_printable_text(file_content)

    # -------------------------------------------------------------------------
    # 3. POWERPOINT PRESENTATIONS (.pptx, .ppt)
    # -------------------------------------------------------------------------
    @staticmethod
    async def extract_text_from_pptx(file_content: bytes) -> str:
        """Extract text slide-by-slide from PPTX/PPT presentations"""
        try:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, DocumentProcessor._process_pptx_sync, file_content)
        except Exception as e:
            logger.warning(f"PPTX extraction fallback: {e}")
            return DocumentProcessor._extract_printable_text(file_content)

    @staticmethod
    def _process_pptx_sync(file_content: bytes) -> str:
        slides_text = []
        try:
            with zipfile.ZipFile(io.BytesIO(file_content)) as z:
                # Find all slide files: ppt/slides/slide1.xml, slide2.xml, etc.
                slide_files = [name for name in z.namelist() if re.match(r'ppt/slides/slide\d+\.xml', name)]
                # Sort numerically by slide number
                slide_files.sort(key=lambda x: int(re.search(r'\d+', x).group()) if re.search(r'\d+', x) else 0)

                for idx, slide_file in enumerate(slide_files, 1):
                    xml_content = z.read(slide_file)
                    tree = ET.fromstring(xml_content)
                    # Find all a:t elements (text runs in PPTX OpenXML)
                    slide_strings = []
                    for elem in tree.iter():
                        if elem.tag.endswith('}t') and elem.text and elem.text.strip():
                            slide_strings.append(elem.text.strip())
                    
                    if slide_strings:
                        title = slide_strings[0] if len(slide_strings) > 0 else f"Slide {idx}"
                        body = "\n• ".join(slide_strings[1:]) if len(slide_strings) > 1 else ""
                        slide_entry = f"--- Slide {idx}: {title} ---\n{body}" if body else f"--- Slide {idx}: {title} ---"
                        slides_text.append(slide_entry)

            if slides_text:
                return "\n\n".join(slides_text).strip()
        except Exception as pptx_err:
            logger.debug(f"PPTX OpenXML extraction error: {pptx_err}")

        return DocumentProcessor._extract_printable_text(file_content)

    # -------------------------------------------------------------------------
    # 4. SPREADSHEETS (.csv, .tsv, .xlsx, .xls)
    # -------------------------------------------------------------------------
    @staticmethod
    async def extract_text_from_spreadsheet(file_content: bytes, filename: str) -> str:
        """Extract text from CSV, TSV, or XLSX files into markdown tables"""
        try:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, DocumentProcessor._process_spreadsheet_sync, file_content, filename)
        except Exception as e:
            logger.warning(f"Spreadsheet extraction fallback: {e}")
            return DocumentProcessor._extract_printable_text(file_content)

    @staticmethod
    def _process_spreadsheet_sync(file_content: bytes, filename: str) -> str:
        ext = filename.lower().split('.')[-1]
        
        # CSV / TSV
        if ext in ['csv', 'tsv']:
            delimiter = '\t' if ext == 'tsv' else ','
            text = DocumentProcessor._decode_text(file_content)
            reader = csv.reader(io.StringIO(text), delimiter=delimiter)
            rows = [row for row in reader if any(cell.strip() for cell in row)]
            if not rows:
                return text
            # Format as markdown table
            header = rows[0]
            divider = ["---"] * len(header)
            md_lines = [
                "| " + " | ".join(header) + " |",
                "| " + " | ".join(divider) + " |"
            ]
            for row in rows[1:100]:  # up to 100 rows
                # Pad row to match header length
                padded = row + [""] * (len(header) - len(row))
                md_lines.append("| " + " | ".join(padded[:len(header)]) + " |")
            return "\n".join(md_lines)

        # XLSX via OpenXML zip
        if ext in ['xlsx', 'xlsm']:
            try:
                with zipfile.ZipFile(io.BytesIO(file_content)) as z:
                    # Shared strings table
                    shared_strings = []
                    if 'xl/sharedStrings.xml' in z.namelist():
                        tree = ET.fromstring(z.read('xl/sharedStrings.xml'))
                        for elem in tree.iter():
                            if elem.tag.endswith('}t') and elem.text:
                                shared_strings.append(elem.text.strip())

                    # Worksheets
                    sheet_files = [n for n in z.namelist() if re.match(r'xl/worksheets/sheet\d+\.xml', n)]
                    sheet_texts = []
                    for s_idx, sheet_file in enumerate(sheet_files, 1):
                        tree = ET.fromstring(z.read(sheet_file))
                        cell_texts = []
                        for cell in tree.iter():
                            if cell.tag.endswith('}v') and cell.text:
                                val = cell.text
                                # Check if it references a shared string
                                if val.isdigit() and int(val) < len(shared_strings):
                                    cell_texts.append(shared_strings[int(val)])
                                else:
                                    cell_texts.append(val)
                        if cell_texts:
                            sheet_texts.append(f"--- Sheet {s_idx} ---\n" + "\n".join(cell_texts))
                    
                    if sheet_texts:
                        return "\n\n".join(sheet_texts).strip()
            except Exception as xlsx_err:
                logger.debug(f"XLSX extraction error: {xlsx_err}")

        return DocumentProcessor._extract_printable_text(file_content)

    # -------------------------------------------------------------------------
    # 5. IMAGES (.png, .jpg, .jpeg, .webp, .bmp, .tiff)
    # -------------------------------------------------------------------------
    @staticmethod
    async def extract_text_from_image(file_content: bytes, filename: str = "image.png") -> str:
        """Extract text from image using OCR with graceful metadata fallback"""
        try:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, DocumentProcessor._process_image_sync, file_content, filename)
        except Exception as e:
            logger.warning(f"Image processing note: {e}")
            return f"Image Study Asset: {filename}\nFormat: Image file uploaded for visual study context."

    @staticmethod
    def _process_image_sync(file_content: bytes, filename: str) -> str:
        # 1. Try Tesseract OCR if available
        if pytesseract and Image:
            try:
                image = Image.open(io.BytesIO(file_content))
                ocr_text = pytesseract.image_to_string(image)
                if ocr_text.strip():
                    return f"--- Scanned Document / OCR Extracted Text: {filename} ---\n\n{ocr_text.strip()}"
            except Exception as ocr_err:
                logger.debug(f"Tesseract OCR failed: {ocr_err}")

        # 2. Extract image metadata and description
        dim_info = ""
        if Image:
            try:
                img = Image.open(io.BytesIO(file_content))
                dim_info = f" ({img.width}x{img.height}, {img.format})"
            except Exception:
                pass

        return f"Image Study Material: {filename}{dim_info}\nVisual note/diagram uploaded as reference material."

    # -------------------------------------------------------------------------
    # 6. TEXT / CODE / MARKDOWN ENCODING UTILITIES
    # -------------------------------------------------------------------------
    @staticmethod
    def _decode_text(file_content: bytes) -> str:
        """Decode text with multi-encoding fallback so it NEVER fails"""
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16', 'gbk', 'shift_jis']
        for enc in encodings:
            try:
                return file_content.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue
        # Fallback with replacement characters
        return file_content.decode('utf-8', errors='replace')

    @staticmethod
    def _extract_printable_text(file_content: bytes) -> str:
        """Extract all printable characters from arbitrary binary data"""
        try:
            decoded = file_content.decode('latin-1', errors='ignore')
            # Extract printable sequences of 3 or more alphanumeric characters
            words = re.findall(r'[A-Za-z0-9\s.,!?:;\-_/\\()\[\]{}@#$%^&*+=<>~"\']{3,}', decoded)
            cleaned = " ".join(words).strip()
            if len(cleaned) > 20:
                return cleaned
        except Exception:
            pass
        return "Study document uploaded successfully."

    # -------------------------------------------------------------------------
    # 7. MAIN UNIVERSAL INGESTION DISPATCHER
    # -------------------------------------------------------------------------
    @staticmethod
    async def process_document(file_content: bytes, filename: str) -> Tuple[str, str]:
        """
        Process document and extract clean text based on file type.
        Supports PDF, DOCX, DOC, PPTX, PPT, CSV, TSV, XLSX, XLS, TXT, MD, IMAGES, CODE.
        NEVER raises an exception on supported or binary formats.
        """
        if not file_content:
            return f"Empty document: {filename}", "text"

        file_extension = filename.lower().split('.')[-1] if '.' in filename else 'txt'

        # PDF
        if file_extension == 'pdf':
            text = await DocumentProcessor.extract_text_from_pdf(file_content)
            return text if text.strip() else f"PDF Document: {filename}", 'pdf'

        # Word Documents
        elif file_extension in ['docx', 'doc', 'dotx', 'dot']:
            text = await DocumentProcessor.extract_text_from_docx(file_content)
            return text if text.strip() else f"Word Document: {filename}", 'docx'

        # PowerPoint Presentations
        elif file_extension in ['pptx', 'ppt', 'ppsx', 'potx']:
            text = await DocumentProcessor.extract_text_from_pptx(file_content)
            return text if text.strip() else f"PowerPoint Presentation: {filename}", 'pptx'

        # Spreadsheets
        elif file_extension in ['csv', 'tsv', 'xlsx', 'xls', 'xlsm']:
            text = await DocumentProcessor.extract_text_from_spreadsheet(file_content, filename)
            return text if text.strip() else f"Spreadsheet Data: {filename}", 'spreadsheet'

        # Images
        elif file_extension in ['png', 'jpg', 'jpeg', 'webp', 'bmp', 'tiff', 'tif', 'gif', 'svg']:
            text = await DocumentProcessor.extract_text_from_image(file_content, filename)
            return text, 'image'

        # Plain Text, Markdown, Notes & Code
        elif file_extension in [
            'txt', 'md', 'markdown', 'rst', 'json', 'py', 'js', 'jsx', 'ts', 'tsx',
            'html', 'htm', 'xml', 'css', 'scss', 'c', 'cpp', 'h', 'hpp', 'java',
            'rs', 'go', 'sql', 'sh', 'bash', 'yaml', 'yml', 'tex', 'latex', 'log'
        ]:
            text = DocumentProcessor._decode_text(file_content)
            return text if text.strip() else f"Text File: {filename}", 'text'

        # Universal fallback for any other file extension
        else:
            text = DocumentProcessor._decode_text(file_content)
            if not text.strip() or len(text.strip()) < 10:
                text = DocumentProcessor._extract_printable_text(file_content)
            return text if text.strip() else f"Uploaded Document: {filename}", 'text'